import PyPDF2
import docx
import io

def test_libs():
    print("Testing PyPDF2...")
    try:
        # Create a dummy PDF in memory
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        c.drawString(100, 750, "Hello World")
        c.save()
        buf.seek(0)
        
        reader = PyPDF2.PdfReader(buf)
        text = reader.pages[0].extract_text()
        print(f"PyPDF2 worked: {text.strip()}")
    except ImportError:
        print("reportlab not installed, skipping PDF generation test")
        # Try just importing and checking version
        print(f"PyPDF2 version: {PyPDF2.__version__}")
    except Exception as e:
        print(f"PyPDF2 Error: {e}")

    print("\nTesting python-docx...")
    try:
        doc = docx.Document()
        doc.add_paragraph("Hello World")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        # Read it back
        doc2 = docx.Document(buf)
        text = doc2.paragraphs[0].text
        print(f"python-docx worked: {text.strip()}")
    except Exception as e:
        print(f"python-docx Error: {e}")

if __name__ == "__main__":
    test_libs()
